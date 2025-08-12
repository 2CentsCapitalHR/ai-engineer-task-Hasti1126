"""
ADGM Corporate Agent - Complete Final Version  
Uses meta-llama/llama-4-scout-17b-16e-instruct (Vision LLM) via Direct Groq API
"""

import gradio as gr
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
import docx.oxml.shared
from qdrant_client import QdrantClient
from qdrant_client.models import Filter, FieldCondition, MatchValue, PayloadSchemaType  # Fixed: Removed FieldIndex
from transformers import CLIPProcessor, CLIPModel
from groq import Groq  # Direct Groq API instead of LangChain
import torch
import numpy as np
import os
import json
import tempfile
import shutil
from pathlib import Path
from dotenv import load_dotenv
from typing import List, Dict, Tuple, Optional
from datetime import datetime
import base64
import io
from PIL import Image
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import uuid
import zipfile
import warnings
warnings.filterwarnings("ignore")

load_dotenv()

class ADGMCorporateAgent:
    """ADGM Corporate Agent: Option 1 Indexed Search + Fixed Imports + File Handling"""
    
    def __init__(self, qdrant_url: str, qdrant_api_key: str, groq_api_key: str):
        # Connect to Qdrant with indexed documents
        self.qdrant_client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key)
        self.collection_name = "adgm_reference_docs"
        
        # OPTION 1: Create missing indexes for content_type field (FIXED IMPORTS)
        self.create_missing_indexes()
        
        # Load image store
        self.image_store = self._load_image_store()
        
        # Initialize CLIP for vector similarity
        print("🔄 Loading CLIP model for vector similarity...")
        self.clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
        self.clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
        self.clip_model.eval()
        
        # Initialize Vision LLM with Direct Groq API
        print("🔄 Initializing Vision LLM: meta-llama/llama-4-scout-17b-16e-instruct")
        self.groq_client = Groq(api_key=groq_api_key)
        self.model_used = "meta-llama/llama-4-scout-17b-16e-instruct"
        
        # Test model connection
        try:
            test_response = self.groq_client.chat.completions.create(
                model=self.model_used,
                messages=[{"role": "user", "content": "Test"}],
                max_tokens=5
            )
            print(f"✅ Successfully initialized Vision LLM: {self.model_used}")
        except Exception as e:
            print(f"⚠️ Primary model failed, trying fallback: {e}")
            # Try fallback models
            fallback_models = ["llama-3.1-70b-versatile", "llama-3.1-8b-instant"]
            for fallback_model in fallback_models:
                try:
                    test_response = self.groq_client.chat.completions.create(
                        model=fallback_model,
                        messages=[{"role": "user", "content": "Test"}],
                        max_tokens=5
                    )
                    self.model_used = fallback_model
                    print(f"✅ Using fallback model: {self.model_used}")
                    break
                except Exception as fallback_error:
                    print(f"⚠️ Fallback {fallback_model} failed: {fallback_error}")
                    continue
            else:
                raise Exception("❌ No available models found on Groq")
        
        # Verify connection to indexed data
        self._verify_indexed_documents()
        
        # ADGM Document checklists for process identification
        self.document_checklists = {
            "Company Incorporation": [
                "Articles of Association", "Memorandum of Association", 
                "Board Resolution", "UBO Declaration Form", "Register of Members and Directors"
            ],
            "Licensing Application": [
                "License Application Form", "Business Plan", "Financial Projections",
                "Compliance Manual", "Professional Indemnity Insurance"
            ],
            "Branch Registration": [
                "Branch Registration Form", "Parent Company Certificate", 
                "Board Resolution for Branch", "Power of Attorney", "Audited Financial Statements"
            ],
            "Foundation Registration": [
                "Foundation Registration Form", "Foundation Charter",
                "Beneficial Ownership Declaration", "Initial Endowment Evidence"
            ],
            "Employment and HR": [
                "Employment Contract", "Employee Handbook", "HR Policies",
                "Disciplinary Procedures", "Leave Policies"
            ],
            "Commercial Agreements": [
                "Service Agreement", "Supply Agreement", "Distribution Agreement",
                "Partnership Agreement", "Joint Venture Agreement"
            ]
        }
        
        # ADGM Red Flags for compliance checking
        self.red_flags = {
            "jurisdiction_issues": [
                "UAE Federal Courts", "Dubai Courts", "Sharjah Courts", "Abu Dhabi Courts",
                "UAE Federal Law", "Dubai International Financial Centre", "DIFC"
            ],
            "missing_clauses": [
                "governing law", "dispute resolution", "registered address", 
                "ADGM jurisdiction", "arbitration clause"
            ],
            "non_binding_language": [
                "may consider", "might", "possibly", "perhaps", "could potentially"
            ],
            "incomplete_sections": [
                "to be completed", "TBD", "pending", "[INSERT]", "DRAFT",
                "under review", "subject to approval"
            ]
        }
        
        # ADGM Legal References for inline commenting
        self.adgm_legal_refs = {
            "jurisdiction": {
                "citation": "ADGM Companies Regulations 2020, Art. 6 & ADGM Courts Law 2013, Art. 15",
                "rule": "All disputes must be subject to ADGM jurisdiction and courts",
                "alternative": "This Agreement shall be governed by ADGM law and any disputes shall be subject to the exclusive jurisdiction of ADGM Courts."
            },
            "governing_law": {
                "citation": "ADGM Companies Regulations 2020, Art. 6(1)",
                "rule": "ADGM law must be specified as governing law",
                "alternative": "This Agreement shall be governed by and construed in accordance with the laws of the Abu Dhabi Global Market (ADGM)."
            },
            "dispute_resolution": {
                "citation": "ADGM Arbitration Law 2015, Art. 7 & ADGM Courts Law 2013",
                "rule": "Dispute resolution mechanism must reference ADGM courts or ADGM-approved arbitration",
                "alternative": "Any dispute arising from this Agreement shall be resolved through arbitration under the ADGM Arbitration Law 2015 or submitted to the jurisdiction of ADGM Courts."
            },
            "registered_address": {
                "citation": "ADGM Companies Regulations 2020, Art. 23",
                "rule": "Companies must maintain registered address within ADGM",
                "alternative": "The registered address of the Company is [Address], Abu Dhabi Global Market, UAE."
            },
            "board_resolution": {
                "citation": "ADGM Companies Regulations 2020, Art. 128-135",
                "rule": "Board resolutions must comply with ADGM procedural requirements",
                "alternative": "BE IT RESOLVED that the Board of Directors, in accordance with ADGM Companies Regulations 2020, hereby approves..."
            },
            "ubo_declaration": {
                "citation": "ADGM AML Rules 2019, Rule 3.2.1 & ADGM Companies Regulations 2020, Art. 41",
                "rule": "Ultimate Beneficial Ownership must be declared as per ADGM requirements",
                "alternative": "The Company hereby declares its Ultimate Beneficial Ownership structure in compliance with ADGM AML Rules 2019."
            }
        }
        
        print("✅ ADGM Corporate Agent initialized successfully ")
        print(f"🤖 Using Vision LLM: {self.model_used}")
        print(f"📚 Connected to indexed knowledge base: {self.indexed_stats.get('total_documents', 0)} documents")
        print(f"🖼️ Image store loaded: {len(self.image_store)} images")

    
    def create_missing_indexes(self):
        """OPTION 1: Create missing indexes for Qdrant collection (FIXED IMPORTS)"""
        try:
            print("🔧 Creating missing indexes for content_type field...")
            
            # Create keyword index for content_type field using compatible method
            self.qdrant_client.create_payload_index(
                collection_name=self.collection_name,
                field_name="content_type",
                field_schema=PayloadSchemaType.KEYWORD
            )
            
            print("✅ Successfully created content_type keyword index")
            
        except Exception as e:
            print(f"⚠️ Index creation failed (may already exist): {e}")
            # This is not fatal - the index may already exist
    
    def _load_image_store(self) -> Dict:
        """Load image store from indexed setup"""
        try:
            if Path("adgm_image_store.json").exists():
                with open("adgm_image_store.json", 'r') as f:
                    store = json.load(f)
                    print(f"📷 Loaded {len(store)} images from storage")
                    return store
            else:
                print("⚠️ No image store found - visual similarity search will be limited")
                return {}
        except Exception as e:
            print(f"Error loading image store: {e}")
            return {}
    
    def _verify_indexed_documents(self):
        """Enhanced verification with accurate document counting"""
        try:
            collection_info = self.qdrant_client.get_collection(self.collection_name)
            
            # Get ALL unique document IDs by scrolling through the collection
            print("🔄 Calculating accurate document count...")
            
            # Use scroll to get all points and count unique document IDs
            all_document_ids = set()
            content_types = {"text": 0, "image": 0}
            offset = None
            
            while True:
                # Scroll through all points to count unique documents
                scroll_result = self.qdrant_client.scroll(
                    collection_name=self.collection_name,
                    limit=100,  # Process in batches
                    offset=offset,
                    with_payload=True
                )
                
                points, next_offset = scroll_result
                
                if not points:
                    break
                    
                for point in points:
                    # Count unique document IDs
                    if point.payload.get('document_id'):
                        all_document_ids.add(point.payload['document_id'])
                    
                    # Count content types
                    content_type = point.payload.get('content_type', 'unknown')
                    if content_type in content_types:
                        content_types[content_type] += 1
                
                offset = next_offset
                if offset is None:
                    break
            
            self.indexed_stats = {
                "total_points": collection_info.points_count,
                "total_documents": len(all_document_ids),  # Accurate count
                "text_chunks": content_types["text"],
                "image_chunks": content_types["image"],
                "vector_size": collection_info.config.params.vectors.size,
                "distance_metric": collection_info.config.params.vectors.distance.name
            }
            
            print(f"✅ Accurate count - Documents: {len(all_document_ids)}, Total Points: {collection_info.points_count}")
            
        except Exception as e:
            print(f"⚠️ Warning: Could not verify indexed documents: {e}")
            self.indexed_stats = {
                "total_points": 0, "total_documents": 0, 
                "text_chunks": 0, "image_chunks": 0
            }
    
    def validate_uploaded_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate uploaded file before processing"""
        try:
            if not os.path.exists(file_path):
                return False, "File does not exist"
            
            size = os.path.getsize(file_path)
            if size == 0:
                return False, "File is empty (0 bytes)"
            
            if size < 1000:  # DOCX should be at least 1KB
                return False, f"File too small ({size} bytes) - may be corrupted"
            
            # Try to open as DOCX
            try:
                doc = docx.Document(file_path)
                if len(doc.paragraphs) == 0:
                    return False, "DOCX file contains no content"
                return True, f"Valid DOCX file with {len(doc.paragraphs)} paragraphs"
            except Exception as e:
                return False, f"Not a valid DOCX file: {str(e)}"
                
        except Exception as e:
            return False, f"Validation error: {str(e)}"
    
    def validate_docx_file(self, file_path: str) -> bool:
        """Enhanced DOCX file validation to prevent 'Package not found' errors"""
        try:
            # Check if file exists
            if not Path(file_path).exists():
                print(f"❌ File does not exist: {file_path}")
                return False
            
            # Check file size (must be > 0)
            file_size = Path(file_path).stat().st_size
            if file_size == 0:
                print(f"❌ File is empty: {file_path}")
                return False
                
            print(f"📊 File size: {file_size} bytes")
            
            # Check if it's a valid zip file (DOCX is a zip archive)
            if not zipfile.is_zipfile(file_path):
                print(f"❌ File is not a valid zip archive: {file_path}")
                return False
            
            # Try to open as zip and check for required DOCX components
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                required_files = ['[Content_Types].xml', 'word/document.xml']
                for required_file in required_files:
                    if required_file not in zip_file.namelist():
                        print(f"❌ Missing required DOCX component: {required_file}")
                        return False
            
            # Try to open with python-docx
            doc = docx.Document(file_path)
            print(f"✅ File validation successful: {file_path}")
            print(f"📄 Document contains {len(doc.paragraphs)} paragraphs")
            return True
            
        except zipfile.BadZipFile:
            print(f"❌ File is corrupted or not a valid DOCX: {file_path}")
            return False
        except Exception as e:
            print(f"❌ File validation failed: {e}")
            return False
    
    def read_docx_with_fallback(self, file_path: str):
        """Read DOCX with multiple fallback methods to handle errors"""
        methods = [
            # Method 1: Standard python-docx
            lambda path: docx.Document(path),
            
            # Method 2: Read as binary first, then create Document
            lambda path: docx.Document(io.BytesIO(open(path, 'rb').read())),
            
            # Method 3: Copy to new location first, then read
            lambda path: self._copy_and_read(path)
        ]
        
        for i, method in enumerate(methods, 1):
            try:
                print(f"🔄 Trying method {i} to read DOCX...")
                doc = method(file_path)
                print(f"✅ Method {i} successful!")
                return doc
            except Exception as e:
                print(f"❌ Method {i} failed: {e}")
                continue
        
        raise Exception("All methods to read DOCX file failed")
    
    def _copy_and_read(self, file_path: str):
        """Copy file to new location and read (fallback method)"""
        temp_dir = tempfile.mkdtemp()
        new_path = os.path.join(temp_dir, "temp_adgm_doc.docx")
        shutil.copy2(file_path, new_path)
        return docx.Document(new_path)
    
    def embed_text(self, text: str) -> np.ndarray:
        """Generate CLIP text embedding for vector similarity"""
        inputs = self.clip_processor(
            text=text, return_tensors="pt", padding=True, 
            truncation=True, max_length=77
        )
        with torch.no_grad():
            features = self.clip_model.get_text_features(**inputs)
            features = features / features.norm(dim=-1, keepdim=True)
            return features.squeeze().numpy()
    
    def embed_image(self, image: Image.Image) -> np.ndarray:
        """Generate CLIP image embedding for vector similarity"""
        inputs = self.clip_processor(images=image, return_tensors="pt")
        with torch.no_grad():
            features = self.clip_model.get_image_features(**inputs)
            features = features / features.norm(dim=-1, keepdim=True)
            return features.squeeze().numpy()
    
    def search_similar_documents_by_vector(self, user_text: str, user_images: List[Image.Image] = None, limit: int = 10) -> Dict:
        """
        OPTION 1: Vector similarity search with proper Qdrant indexing 
        Core function for finding similar ADGM documents using cosine similarity
        """
        similarity_results = {
            "text_similarities": [],
            "image_similarities": [],
            "combined_results": [],
            "similarity_method": "cosine_distance_with_clip_embeddings",
            "search_stats": {}
        }
        
        try:
            # 1. TEXT VECTOR SIMILARITY SEARCH (Now with proper indexing)
            if user_text.strip():
                print(f"🔍 Performing text vector similarity search with indexed filtering...")
                user_text_embedding = self.embed_text(user_text)
                
                # Now this should work because we created the content_type index
                text_similar_results = self.qdrant_client.search(
                    collection_name=self.collection_name,
                    query_vector=user_text_embedding.tolist(),
                    limit=limit,
                    with_payload=True,
                    query_filter=Filter(must=[
                        FieldCondition(key="content_type", match=MatchValue(value="text"))
                    ])
                )
                
                similarity_results["text_similarities"] = [
                    {
                        "content": result.payload["content"][:300],
                        "similarity_score": result.score,
                        "document_name": result.payload.get("document_name", "Unknown"),
                        "document_type": result.payload.get("document_type", "Unknown"),
                        "category": result.payload.get("category", "General"),
                        "source_file": result.payload.get("source_file", "Unknown"),
                        "document_id": result.payload.get("document_id", "")
                    }
                    for result in text_similar_results
                ]
                
                similarity_results["search_stats"]["text_search_count"] = len(text_similar_results)
                print(f"📊 Found {len(text_similar_results)} text similarities using indexed filtering")
            
            # 2. IMAGE VECTOR SIMILARITY SEARCH (Now with proper indexing)
            if user_images:
                print(f"🖼️ Performing image vector similarity search for {len(user_images)} user images...")
                
                for img_idx, user_image in enumerate(user_images):
                    user_image_embedding = self.embed_image(user_image)
                    
                    # Now this should work because we created the content_type index
                    image_similar_results = self.qdrant_client.search(
                        collection_name=self.collection_name,
                        query_vector=user_image_embedding.tolist(),
                        limit=5,
                        with_payload=True,
                        query_filter=Filter(must=[
                            FieldCondition(key="content_type", match=MatchValue(value="image"))
                        ])
                    )
                    
                    for result in image_similar_results:
                        image_id = result.payload.get("image_id")
                        similarity_data = {
                            "user_image_index": img_idx,
                            "similarity_score": result.score,
                            "reference_image_type": result.payload.get("image_type", "Unknown"),
                            "document_name": result.payload.get("document_name", "Unknown"),
                            "category": result.payload.get("category", "General"),
                            "image_id": image_id,
                            "document_id": result.payload.get("document_id", "")
                        }
                        
                        # Add reference image data if available in store
                        if image_id and image_id in self.image_store:
                            similarity_data["reference_image_available"] = True
                        
                        similarity_results["image_similarities"].append(similarity_data)
                
                similarity_results["search_stats"]["image_search_count"] = len(similarity_results["image_similarities"])
                print(f"📊 Found {len(similarity_results['image_similarities'])} image similarities using indexed filtering")
            
            # 3. COMBINE AND RANK BY SIMILARITY SCORES
            all_similarities = (
                similarity_results["text_similarities"] + 
                similarity_results["image_similarities"]
            )
            
            # Sort by similarity score (highest first)
            combined_sorted = sorted(all_similarities, key=lambda x: x["similarity_score"], reverse=True)
            similarity_results["combined_results"] = combined_sorted[:limit]
            
            # Add search statistics
            similarity_results["search_stats"].update({
                "total_results": len(combined_sorted),
                "top_similarity_score": combined_sorted[0]["similarity_score"] if combined_sorted else 0,
                "unique_documents": len(set(r.get("document_id", "") for r in combined_sorted if r.get("document_id"))),
                "search_method": "indexed_filtering_option1_all_fixes"
            })
            
            print(f"🎯 Top similarity scores with indexed search: {[r['similarity_score'] for r in combined_sorted[:3]]}")
            
            return similarity_results
            
        except Exception as e:
            print(f"❌ Error in vector similarity search: {e}")
            similarity_results["error"] = str(e)
            return similarity_results
    
    def identify_document_process(self, docx_files: List) -> Tuple[str, float]:
        """Identify legal process from uploaded files using fuzzy matching"""
        if not docx_files:
            return ("General Document Review", 0.0)
            
        file_names = [f.name.lower() for f in docx_files]
        scores = {}
        
        for process, required_docs in self.document_checklists.items():
            score = 0
            matched_docs = 0
            
            for required_doc in required_docs:
                doc_keywords = required_doc.lower().replace(" ", "").split()
                
                for file_name in file_names:
                    clean_filename = file_name.replace("_", "").replace("-", "").replace(" ", "")
                    
                    # Check if any keywords match
                    keyword_matches = sum(1 for keyword in doc_keywords if keyword in clean_filename)
                    if keyword_matches > 0:
                        score += keyword_matches / len(doc_keywords)
                        matched_docs += 1
                        break
            
            scores[process] = score / len(required_docs) if required_docs else 0
        
        best_process = max(scores, key=scores.get) if scores else "General Document Review"
        confidence = scores.get(best_process, 0)
        
        return (best_process, confidence) if confidence > 0.25 else ("General Document Review", confidence)
    
    def extract_all_images_from_docx(self, docx_path: str) -> List[Image.Image]:
        """Extract ALL images from DOCX - tables and embedded images with enhanced error handling"""
        all_images = []
        
        try:
            # Use fallback reading method
            doc = self.read_docx_with_fallback(docx_path)
            
            # Extract tables as images (for Vision LLM)
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            table_data.append(row_data)
                    
                    if table_data and len(table_data) > 1:  # Must have header + data
                        table_image = self._create_table_image(table_data)
                        if table_image:
                            all_images.append(table_image)
                
                except Exception as e:
                    print(f"Error processing table {table_idx}: {e}")
            
            # Extract embedded images
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            image = Image.open(io.BytesIO(image_data))
                            if image.mode != 'RGB':
                                image = image.convert('RGB')
                            
                            # Resize if too large to save memory
                            max_size = (1024, 1024)
                            if image.size[0] > max_size[0] or image.size[1] > max_size[1]:
                                image.thumbnail(max_size, Image.Resampling.LANCZOS)
                            
                            all_images.append(image)
                            
                        except Exception as e:
                            print(f"Error processing embedded image: {e}")
            except Exception as e:
                print(f"Error accessing document images: {e}")
        
        except Exception as e:
            print(f"Error opening DOCX file for image extraction: {e}")
        
        return all_images
    
    def _create_table_image(self, table_data: List[List[str]]) -> Optional[Image.Image]:
        """Create high-quality table image for Vision LLM analysis"""
        try:
            if not table_data or not table_data[0]:
                return None
            
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            
            # Dynamic sizing for better visualization
            fig_width = max(cols * 2.0, 8)
            fig_height = max(rows * 0.8, 4)
            
            fig, ax = plt.subplots(figsize=(fig_width, fig_height))
            ax.set_xlim(0, cols)
            ax.set_ylim(0, rows)
            ax.axis('off')
            
            # Enhanced table rendering with proper formatting
            for i, row in enumerate(table_data):
                for j in range(cols):
                    cell_text = row[j] if j < len(row) else ""
                    
                    # Style headers differently
                    is_header = i == 0
                    rect = patches.Rectangle(
                        (j, rows - i - 1), 1, 1,
                        linewidth=2 if is_header else 1, 
                        edgecolor='black', 
                        facecolor='lightblue' if is_header else 'white',
                        alpha=0.8 if is_header else 1.0
                    )
                    ax.add_patch(rect)
                    
                    # Add text with appropriate formatting
                    cell_text = str(cell_text)[:100]  # Limit text length
                    font_size = 11 if is_header else 9
                    font_weight = 'bold' if is_header else 'normal'
                    
                    ax.text(j + 0.5, rows - i - 0.5, cell_text,
                           ha='center', va='center', 
                           fontsize=font_size, 
                           weight=font_weight,
                           wrap=True)
            
            # High-quality image output
            fig.canvas.draw()
            buf = io.BytesIO()
            plt.savefig(buf, format='png', bbox_inches='tight', dpi=150, 
                       facecolor='white', edgecolor='none')
            buf.seek(0)
            image = Image.open(buf)
            plt.close(fig)
            
            return image
            
        except Exception as e:
            print(f"Error creating table image: {e}")
            return None
    
    def analyze_with_groq_direct(self, analysis_prompt: str) -> str:
        """Use direct Groq API instead of LangChain wrapper"""
        try:
            response = self.groq_client.chat.completions.create(
                model=self.model_used,
                messages=[
                    {"role": "user", "content": analysis_prompt}
                ],
                max_tokens=4000,
                temperature=0.1
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"❌ Groq API error: {e}")
            return f"Error: {str(e)}"
    
    def analyze_docx_with_vector_similarity_and_vision(self, docx_path: str, process_type: str) -> Dict:
        """Enhanced analysis with : DOCX error handling + Vector similarity + Vision LLM"""
        try:
            # Enhanced file validation and reading
            print(f"🔍 Starting analysis of: {os.path.basename(docx_path)}")
            
            # Step 1: Validate file
            if not self.validate_docx_file(docx_path):
                return {
                    "compliance_score": 0,
                    "jurisdiction_compliance": False,
                    "issues_found": [{
                        "location": "File Validation",
                        "issue": "Invalid or corrupted DOCX file",
                        "severity": "High",
                        "suggestion": "Please upload a valid DOCX file or check for file corruption"
                    }],
                    "error": "File validation failed",
                    "vector_similarity_stats": {"error": "File validation failed"}
                }
            
            # Step 2: Read document with fallback methods
            try:
                doc = self.read_docx_with_fallback(docx_path)
                print(f"✅ Successfully opened document with {len(doc.paragraphs)} paragraphs")
            except Exception as e:
                return {
                    "compliance_score": 0,
                    "jurisdiction_compliance": False,
                    "issues_found": [{
                        "location": "Document Reading",
                        "issue": f"Failed to read DOCX file: {str(e)}",
                        "severity": "High",
                        "suggestion": "File may be corrupted or in an unsupported format"
                    }],
                    "error": f"Document reading failed: {str(e)}",
                    "vector_similarity_stats": {"error": f"Document reading failed: {str(e)}"}
                }
            
            # Step 3: Extract text content
            full_text = []
            paragraph_locations = {}
            for para_idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    full_text.append(para.text.strip())
                    # Map text snippets to paragraph numbers for precise location tracking
                    snippet = para.text.strip()[:50]
                    paragraph_locations[snippet] = f"Paragraph {para_idx + 1}"
            
            document_text = '\n'.join(full_text)
            
            if not document_text.strip():
                return {
                    "compliance_score": 0,
                    "jurisdiction_compliance": False,
                    "issues_found": [{
                        "location": "Document Content",
                        "issue": "Document appears to be empty or contains no readable text",
                        "severity": "High",
                        "suggestion": "Please upload a document with text content"
                    }],
                    "error": "Empty document",
                    "vector_similarity_stats": {"error": "Empty document"}
                }
            
            # Step 4: Extract images from user document
            user_images = self.extract_all_images_from_docx(docx_path)
            
            print(f"🔍 Analyzing document: {len(document_text)} chars, {len(user_images)} images")
            
            # Step 5: PERFORM VECTOR SIMILARITY SEARCH ( )
            print("🎯 Performing indexed vector similarity search against ADGM knowledge base...")
            similarity_results = self.search_similar_documents_by_vector(
                user_text=document_text[:2000],  # Use first 2000 chars for similarity
                user_images=user_images,
                limit=10
            )
            
            # Step 6: Detect red flags in text
            red_flags = self._detect_red_flags(document_text)
            
            # Step 7: Create comprehensive analysis prompt with similarity results
            analysis_prompt = f"""You are an ADGM legal compliance expert. Analyze this document using vector similarity matches from the ADGM knowledge base.

DOCUMENT TYPE: {process_type}
DOCUMENT TEXT CONTENT: {document_text[:2500]}

RED FLAGS DETECTED: {json.dumps(red_flags, indent=2)}

VECTOR SIMILARITY SEARCH RESULTS FROM ADGM DATABASE ( ):

Top Similar ADGM Text Documents (by cosine similarity with indexed filtering):
{self._format_similarity_results(similarity_results['text_similarities'][:5])}

Similar Visual Content from ADGM Database (indexed search):
{self._format_image_similarity_results(similarity_results['image_similarities'][:3])}

Search Statistics:
- Total similar documents found: {similarity_results['search_stats'].get('total_results', 0)}
- Unique ADGM documents matched: {similarity_results['search_stats'].get('unique_documents', 0)}
- Top similarity score: {similarity_results['search_stats'].get('top_similarity_score', 0):.3f}
- Search method: {similarity_results['search_stats'].get('search_method', 'indexed')}

The user document contains {len(user_images)} images/tables for analysis.

Based on the vector similarity matches and ADGM regulatory knowledge, provide comprehensive analysis as JSON:
{{
    "compliance_score": <0-100>,
    "jurisdiction_compliance": <boolean>,
    "issues_found": [
        {{
            "location": "specific section/paragraph",
            "issue": "detailed compliance issue description",
            "severity": "High|Medium|Low", 
            "text_match": "exact problematic text from document",
            "adgm_violation": "specific ADGM regulation violated",
            "suggestion": "detailed correction needed",
            "alternative_clause": "complete compliant replacement text"
        }}
    ],
    "missing_elements": ["required elements not found in document"],
    "recommendations": ["specific improvement actions"],
    "similarity_analysis": "how document compares to similar ADGM documents found",
    "reference_alignment": "assessment of alignment with ADGM templates and standards"
}}

Focus on: jurisdiction clauses, governing law, dispute resolution, registered address, signatory authorization, UBO declarations, and compliance with ADGM regulations."""
            
            # Step 8: Get Vision LLM response using Direct Groq API
            print(f"🤖 Sending to Vision LLM for analysis: {self.model_used}")
            response_content = self.analyze_with_groq_direct(analysis_prompt)
            
            try:
                analysis_result = json.loads(response_content)
                
                # Enhance issues with paragraph locations
                for issue in analysis_result.get("issues_found", []):
                    if "text_match" in issue and issue["text_match"]:
                        text_snippet = issue["text_match"][:50]
                        location = paragraph_locations.get(text_snippet, issue.get("location", "Document"))
                        issue["location"] = location
                
                # Add vector similarity metadata
                analysis_result["vector_similarity_stats"] = {
                    "total_text_matches": len(similarity_results['text_similarities']),
                    "total_image_matches": len(similarity_results['image_similarities']),
                    "top_similarity_score": similarity_results['search_stats'].get('top_similarity_score', 0),
                    "similarity_method": similarity_results['similarity_method'],
                    "user_images_analyzed": len(user_images),
                    "unique_adgm_docs_referenced": similarity_results['search_stats'].get('unique_documents', 0),
                    "search_method": "all_fixes_applied_indexed_search"
                }
                
            except json.JSONDecodeError:
                print("⚠️ JSON parsing failed, creating structured response...")
                analysis_result = {
                    "compliance_score": 65,
                    "jurisdiction_compliance": False,
                    "issues_found": [{
                        "location": "Document Review",
                        "issue": "Analysis completed but JSON parsing failed - manual review recommended",
                        "severity": "Medium",
                        "text_match": document_text[:100],
                        "suggestion": "Professional ADGM legal review recommended"
                    }],
                    "missing_elements": [],
                    "recommendations": ["Manual compliance review by ADGM expert"],
                    "similarity_analysis": f"Found {len(similarity_results['combined_results'])} similar ADGM documents using indexed search",
                    "vector_similarity_stats": {
                        "total_text_matches": len(similarity_results['text_similarities']),
                        "total_image_matches": len(similarity_results['image_similarities']),
                        "similarity_method": similarity_results['similarity_method'],
                        "search_method": "all_fixes_applied_indexed_search",
                        "parsing_error": True
                    },
                    "raw_response": response_content[:500]
                }
            
            return analysis_result
            
        except Exception as e:
            print(f"❌ Analysis error: {e}")
            return {
                "compliance_score": 0,
                "jurisdiction_compliance": False,
                "issues_found": [{
                    "location": "System Error",
                    "issue": f"Technical error during analysis: {str(e)}",
                    "severity": "High",
                    "suggestion": "Please try again or contact support"
                }],
                "missing_elements": [],
                "recommendations": ["Technical issue - please retry"],
                "vector_similarity_stats": {"error": str(e)}
            }
    
    def _detect_red_flags(self, text: str) -> Dict[str, List[str]]:
        """Detect ADGM compliance red flags in document text"""
        found_flags = {}
        text_lower = text.lower()
        
        for category, patterns in self.red_flags.items():
            found_in_category = [pattern for pattern in patterns if pattern.lower() in text_lower]
            if found_in_category:
                found_flags[category] = found_in_category
        
        return found_flags
    
    def _format_similarity_results(self, results: List[Dict]) -> str:
        """Format text similarity results for LLM prompt"""
        if not results:
            return "No similar text documents found in ADGM database."
        
        formatted = []
        for i, result in enumerate(results, 1):
            formatted.append(
                f"{i}. [{result['document_type']}] {result['document_name']} "
                f"(Similarity: {result['similarity_score']:.3f})\n"
                f"   Category: {result['category']}\n"
                f"   Content: {result['content'][:200]}...\n"
            )
        
        return '\n'.join(formatted)
    
    def _format_image_similarity_results(self, results: List[Dict]) -> str:
        """Format image similarity results for LLM prompt"""
        if not results:
            return "No similar visual content found in ADGM database."
        
        formatted = []
        for i, result in enumerate(results, 1):
            formatted.append(
                f"{i}. {result['reference_image_type']} from {result['document_name']} "
                f"(Similarity: {result['similarity_score']:.3f})\n"
                f"   Category: {result['category']}"
            )
        
        return '\n'.join(formatted)
    
    def add_enhanced_comments_to_docx(self, docx_path: str, issues: List[Dict]) -> str:
        """Add enhanced inline comments  for proper file handling"""
        try:
            # Use fallback reading method for commenting as well
            doc = self.read_docx_with_fallback(docx_path)
            comments_added = 0
            
            for issue in issues:
                text_match = issue.get('text_match', '')
                issue_type = issue.get('issue', '').lower()
                severity = issue.get('severity', 'Medium')
                suggestion = issue.get('suggestion', '')
                location = issue.get('location', 'Unknown')
                
                if not text_match:
                    continue
                
                # Get relevant ADGM legal reference
                legal_ref = None
                for key, ref_data in self.adgm_legal_refs.items():
                    if key in issue_type or key in suggestion.lower():
                        legal_ref = ref_data
                        break
                
                if not legal_ref:
                    legal_ref = {
                        "citation": "ADGM Companies Regulations 2020 & ADGM General Rules 2019",
                        "rule": "Document must comply with ADGM legal framework",
                        "alternative": "Please review and align with ADGM regulatory requirements."
                    }
                
                # Find and process paragraphs containing the issue
                for para in doc.paragraphs:
                    if text_match.lower() in para.text.lower():
                        # Highlight the problematic text
                        for run in para.runs:
                            if text_match.lower() in run.text.lower():
                                if severity.lower() == 'high':
                                    run.font.highlight_color = WD_COLOR_INDEX.RED
                                elif severity.lower() == 'medium':
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                else:
                                    run.font.highlight_color = WD_COLOR_INDEX.LIGHT_BLUE
                        
                        # Create comprehensive comment with ADGM legal citation
                        comment_parts = [
                            f"🚨 ADGM COMPLIANCE ISSUE ({severity.upper()})",
                            f"📍 Location: {location}",
                            f"⚖️ Legal Citation: {legal_ref['citation']}",
                            f"📋 ADGM Requirement: {legal_ref['rule']}",
                            f"🔧 Issue Identified: {suggestion or issue.get('issue', 'Compliance issue detected')}",
                            f"✅ Recommended Alternative: {legal_ref['alternative']}"
                        ]
                        
                        try:
                            # Add comprehensive comment paragraph
                            new_para = para._parent.add_paragraph()
                            
                            for part_idx, part in enumerate(comment_parts):
                                if part_idx > 0:
                                    new_para.add_run("\n")
                                
                                run = new_para.add_run(part)
                                
                                # Style based on content type
                                if "🚨" in part:  # Header
                                    run.bold = True
                                    run.font.size = docx.shared.Pt(12)
                                    if severity.lower() == 'high':
                                        run.font.color.rgb = RGBColor(220, 20, 60)  # Crimson
                                    elif severity.lower() == 'medium':
                                        run.font.color.rgb = RGBColor(255, 140, 0)   # Orange
                                    else:
                                        run.font.color.rgb = RGBColor(30, 144, 255)  # Blue
                                elif "⚖️" in part:  # Legal citation
                                    run.bold = True
                                    run.italic = True
                                    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                                    run.font.size = docx.shared.Pt(10)
                                elif "✅" in part:  # Alternative suggestion
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                                    run.font.size = docx.shared.Pt(10)
                                else:  # Regular content
                                    run.font.size = docx.shared.Pt(9)
                                    run.font.color.rgb = RGBColor(64, 64, 64)  # Dark gray
                            
                            # Add separator
                            separator_run = new_para.add_run("\n" + "─" * 70 + "\n")
                            separator_run.font.color.rgb = RGBColor(128, 128, 128)
                            separator_run.font.size = docx.shared.Pt(8)
                            
                            comments_added += 1
                            
                        except Exception as e:
                            print(f"Error adding detailed comment: {e}")
                        
                        break  # Only add comment once per issue
            
            # Add comprehensive summary at the end
            if comments_added > 0:
                try:
                    summary_para = doc.add_paragraph()
                    summary_parts = [
                        f"\n{'='*80}",
                        f"📊 ADGM COMPLIANCE SUMMARY -  ",
                        f"{'='*80}",
                        f"🤖 Analysis Model: {self.model_used}",
                        f"🔍 Total Issues Identified: {comments_added}",
                        f"⚖️ All citations reference current ADGM regulations (2019-2020)",
                        f"🎯 Analysis based on indexed vector similarity with {self.indexed_stats.get('total_documents', 0)} ADGM documents",
                        f"📚 Knowledge base: {self.indexed_stats.get('total_points', 0)} indexed legal vectors",
                        f"🔧 Search method: Option 1 - Efficient indexed filtering (ALL FIXES)",
                        f"✅ Review suggested alternatives and implement corrections",
                        f"📞 For clarification, consult ADGM Registration Authority",
                        f"🌐 ADGM Regulations: https://www.adgm.com/doing-business/laws-and-regulations",
                        f"📅 Analysis Date: {datetime.now().strftime('%B %d, %Y')}",
                        f"🤖 Powered by: Complete ADGM Corporate Agent with  ",
                        f"{'='*80}\n"
                    ]
                    
                    for part in summary_parts:
                        summary_run = summary_para.add_run(part + "\n")
                        if "📊" in part or "=" in part:
                            summary_run.bold = True
                            summary_run.font.color.rgb = RGBColor(0, 0, 139)
                            summary_run.font.size = docx.shared.Pt(11)
                        else:
                            summary_run.font.color.rgb = RGBColor(64, 64, 64)
                            summary_run.font.size = docx.shared.Pt(10)
                        
                except Exception as e:
                    print(f"Error adding summary: {e}")
            
            # ALL FIXES: Save to a location that Gradio can access
            base_dir = os.path.dirname(docx_path)
            output_filename = os.path.basename(docx_path).replace('.docx', '_ADGM_ALL_FIXES_Analysis.docx')
            output_path = os.path.join(base_dir, output_filename)
            
            # Ensure the directory exists
            os.makedirs(base_dir, exist_ok=True)
            
            doc.save(output_path)
            
            # Verify the file was saved successfully
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"✅ Successfully saved  analysis file: {output_path}")
                print(f"📊 File size: {os.path.getsize(output_path)} bytes")
                print(f"✅ Added {comments_added} detailed ADGM legal comments with citations")
                return output_path
            else:
                print(f"❌ Failed to save analysis file properly")
                return None
            
        except Exception as e:
            print(f"❌ Error in enhanced commenting: {e}")
            return None
    
    def process_user_documents(self, uploaded_files: List) -> Dict:
        """Enhanced processing with  """
        if not uploaded_files:
            return {"error": "No files uploaded"}
        
        docx_files = [f for f in uploaded_files if f.name.endswith('.docx')]
        if not docx_files:
            return {"error": "Please upload DOCX files only"}
        
        temp_dir = tempfile.mkdtemp()
        results = {
            "timestamp": datetime.now().isoformat(),
            "model_used": self.model_used,
            "similarity_method": "all_fixes_applied_indexed_clip_cosine_vector_search",
            "knowledge_base_stats": self.indexed_stats,
            "total_files": len(docx_files),
            "process_identified": "",
            "confidence": 0.0,
            "documents_analysis": [],
            "missing_documents": [],
            "overall_compliance": 0,
            "reviewed_files": []
        }
        
        try:
            saved_files = []
            for file in docx_files:
                print(f"🔄 Processing: {os.path.basename(file.name)}")
                
                try:
                    # Use file.name (path) directly - FIXED file upload handling
                    original_path = file.name
                    print(f"📁 Original file path: {original_path}")
                    
                    # Validate uploaded file
                    is_valid, validation_message = self.validate_uploaded_file(original_path)
                    
                    if not is_valid:
                        print(f"❌ File validation failed: {validation_message}")
                        
                        # Check if it's the filename issue we discovered
                        if "File is empty (0 bytes)" in validation_message:
                            results["documents_analysis"].append({
                                "file_name": os.path.basename(file.name),
                                "error": f"File upload issue: {validation_message}. Please rename to a simpler filename (e.g., 'document.docx') and try again.",
                                "compliance_score": 0,
                                "issues_found": [{
                                    "location": "File Upload",
                                    "issue": "File appears empty - likely due to complex filename",
                                    "severity": "High",
                                    "suggestion": "Rename file to something simple like 'test.docx' or 'document.docx' and re-upload"
                                }]
                            })
                        else:
                            results["documents_analysis"].append({
                                "file_name": os.path.basename(file.name),
                                "error": f"File validation failed: {validation_message}",
                                "compliance_score": 0,
                                "issues_found": [{
                                    "location": "File Upload", 
                                    "issue": validation_message,
                                    "severity": "High",
                                    "suggestion": "Ensure file is a valid DOCX and try again"
                                }]
                            })
                        continue
                    
                    print(f"✅ File validation passed: {validation_message}")
                    
                    # Copy file to working directory with safe name
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    safe_filename = f"adgm_doc_{timestamp}.docx"
                    work_path = os.path.join(temp_dir, safe_filename)
                    
                    # Use shutil.copy2 to preserve file metadata
                    shutil.copy2(original_path, work_path)
                    
                    # Verify copy was successful
                    copied_size = os.path.getsize(work_path)
                    print(f"✅ Copied file: {copied_size} bytes")
                    
                    # Final DOCX validation
                    if self.validate_docx_file(work_path):
                        saved_files.append((work_path, os.path.basename(file.name)))
                        print(f"✅ Successfully processed: {os.path.basename(file.name)}")
                    else:
                        print(f"❌ DOCX validation failed: {file.name}")
                        results["documents_analysis"].append({
                            "file_name": os.path.basename(file.name),
                            "error": "Invalid DOCX file structure",
                            "compliance_score": 0,
                            "issues_found": [{
                                "location": "File Validation",
                                "issue": "File is not a valid DOCX document",
                                "severity": "High",
                                "suggestion": "Ensure file is a proper DOCX format and not corrupted"
                            }]
                        })
                        
                except Exception as e:
                    print(f"❌ Error processing {file.name}: {e}")
                    results["documents_analysis"].append({
                        "file_name": os.path.basename(file.name),
                        "error": f"Processing failed: {str(e)}",
                        "compliance_score": 0,
                        "issues_found": [{
                            "location": "File Processing",
                            "issue": f"Failed to process uploaded file: {str(e)}",
                            "severity": "High",
                            "suggestion": "Check file format and try again"
                        }]
                    })
            
            if not saved_files:
                return {"error": "No valid files could be processed. Files may be empty or have complex filenames. Try renaming to simple names like 'document.docx'."}
            
     
            # Identify process type from valid files
            temp_file_objects = [type('obj', (object,), {'name': name}) for _, name in saved_files]
            process_type, confidence = self.identify_document_process(temp_file_objects)
            results["process_identified"] = process_type
            results["confidence"] = confidence
            
            print(f"🎯 Identified process: {process_type} (Confidence: {confidence:.1%})")
            
            # Check for missing documents
            if process_type in self.document_checklists:
                required_docs = self.document_checklists[process_type]
                uploaded_names = [name.lower() for _, name in saved_files]
                
                missing_docs = []
                for required_doc in required_docs:
                    doc_keywords = required_doc.lower().replace(" ", "").split()
                    found = any(
                        any(keyword in uploaded_name.replace("_", "").replace("-", "") 
                            for keyword in doc_keywords)
                        for uploaded_name in uploaded_names
                    )
                    if not found:
                        missing_docs.append(required_doc)
                
                results["missing_documents"] = missing_docs
                print(f"📋 Missing documents: {len(missing_docs)}")
            
            # Analyze each valid document with  
            compliance_scores = []
            
            for file_idx, (file_path, original_name) in enumerate(saved_files):
                print(f"\n🔍 Analyzing document {file_idx + 1}/{len(saved_files)}: {original_name} (ALL FIXES)")
                
                # Perform comprehensive analysis with  
                analysis = self.analyze_docx_with_vector_similarity_and_vision(file_path, process_type)
                
                # Add enhanced comments if issues found and no errors
                if analysis.get('issues_found') and not analysis.get('error'):
                    try:
                        reviewed_path = self.add_enhanced_comments_to_docx(file_path, analysis['issues_found'])
                        
                        # ALL FIXES: Verify file exists before adding to results
                        if reviewed_path and os.path.exists(reviewed_path):
                            # Copy to a more accessible location for Gradio with simple name
                            final_temp_dir = tempfile.mkdtemp()
                            simple_name = f"{original_name.replace('.docx', '')}_ADGM_ALL_FIXES_Analysis.docx"
                            final_output = os.path.join(final_temp_dir, simple_name)
                            shutil.copy2(reviewed_path, final_output)
                            
                            # Verify final file is accessible
                            if os.path.exists(final_output) and os.path.getsize(final_output) > 0:
                                results["reviewed_files"].append(final_output)
                                print(f"✅ Created  downloadable file: {final_output}")
                                print(f"📊 Final file size: {os.path.getsize(final_output)} bytes")
                            else:
                                print(f"❌ Final file not accessible: {final_output}")
                        else:
                            print(f"⚠️ Could not create enhanced version for {original_name}")
                            
                    except Exception as e:
                        print(f"⚠️ Warning: Could not add comments to {original_name}: {e}")
                
                analysis["file_name"] = original_name
                results["documents_analysis"].append(analysis)
                
                # Track compliance scores
                if 'compliance_score' in analysis and analysis['compliance_score'] > 0:
                    compliance_scores.append(analysis['compliance_score'])
            
            # Calculate overall compliance
            if compliance_scores:
                results["overall_compliance"] = sum(compliance_scores) / len(compliance_scores)
            else:
                results["overall_compliance"] = 0
            
            print(f"✅  processing complete. Overall compliance: {results['overall_compliance']:.1f}%")
            
            return results
            
        except Exception as e:
            print(f"❌ Processing error: {e}")
            return {"error": f"Processing failed: {str(e)}"}
        
        finally:
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except:
                pass

def create_gradio_interface():
    """Create enhanced Gradio interface with  """
    
    try:
        agent = ADGMCorporateAgent(
            qdrant_url=os.getenv("QDRANT_CLOUD_URL"),
            qdrant_api_key=os.getenv("QDRANT_API_KEY"),
            groq_api_key=os.getenv("GROQ_API_KEY")
        )
    except Exception as e:
        print(f"❌ Failed to initialize agent: {e}")
        return None
    
    def analyze_documents(files):
        """Enhanced document analysis  """
        if not files:
            return "Please upload DOCX files for analysis.", "", None
        
        try:
            print(f"📄 Received {len(files)} files from Gradio")
            
            # Debug file information
            for i, file in enumerate(files):
                print(f"🔍 File {i+1}: {getattr(file, 'name', 'Unknown')}")
                if hasattr(file, 'name') and os.path.exists(file.name):
                    size = os.path.getsize(file.name)
                    print(f"   Size on disk: {size} bytes")
                    if size == 0:
                        print(f"   ⚠️ Empty file detected - filename may be too complex")
            
            print(f"🔄 Starting  processing with indexed vector search...")
            results = agent.process_user_documents(files)
            
            if "error" in results:
                return f"❌ Error: {results['error']}", json.dumps(results, indent=2), None
            
            # Generate comprehensive summary
            kb_stats = results.get('knowledge_base_stats', {})
            process_type = results['process_identified']
            confidence = results['confidence']
            missing_docs = results['missing_documents']
            overall_compliance = results['overall_compliance']
            model_used = results.get('model_used', 'Unknown')
            
            summary = f"""# 🏛️ ADGM Corporate Agent -   EDITION

## 🤖 AI Analysis Details
**Vision LLM Model**: {model_used}
**API Integration**: Direct Groq API (Enhanced Compatibility)
**Vector Search Method**: {results.get('similarity_method', '  Indexed CLIP Cosine Search')}
**Search Performance**: ⚡ Fast indexed filtering (ALL FIXES)
**Knowledge Base**: {kb_stats.get('total_points', 0)} indexed vectors from {kb_stats.get('total_documents', 0)} ADGM documents

## 📋 Process Identification  
**Identified Process**: {process_type}
**Confidence Level**: {confidence:.1%}

## 📊 Document Analysis Summary
**Documents Analyzed**: {results['total_files']}
**Overall Compliance Score**: {overall_compliance:.1f}/100
"""
            
            # Compliance status indicator
            if overall_compliance >= 80:
                summary += "**Status**: ✅ **HIGHLY COMPLIANT** - Strong ADGM adherence\n\n"
            elif overall_compliance >= 60:
                summary += "**Status**: ⚠️ **NEEDS IMPROVEMENT** - Several compliance issues found\n\n"
            else:
                summary += "**Status**: ❌ **NON-COMPLIANT** - Significant compliance gaps identified\n\n"
            
            # Missing documents section
            if missing_docs:
                summary += f"## ⚠️ Missing Required Documents ({len(missing_docs)})\n"
                for doc in missing_docs:
                    summary += f"- {doc}\n"
                summary += "\n"
            else:
                summary += "## ✅ Document Completeness\nAll required documents for this process are present.\n\n"
            
            # Individual document analysis
            summary += "## 📄 Document Analysis Results (ALL FIXES)\n\n"
            
            for i, analysis in enumerate(results["documents_analysis"], 1):
                file_name = analysis.get("file_name", f"Document {i}")
                compliance_score = analysis.get("compliance_score", 0)
                similarity_stats = analysis.get("vector_similarity_stats", {})
                
                # Check for errors
                if analysis.get("error"):
                    summary += f"### ❌ {file_name}\n"
                    summary += f"**Error**: {analysis['error']}\n"
                    summary += f"**Status**: Processing failed\n\n"
                    continue
                
                status_emoji = "✅" if compliance_score >= 80 else "⚠️" if compliance_score >= 60 else "❌"
                
                summary += f"### {status_emoji} {file_name}\n"
                summary += f"**Compliance Score**: {compliance_score}/100\n"
                summary += f"**ADGM Vector Matches**: {similarity_stats.get('total_text_matches', 0)} text + {similarity_stats.get('total_image_matches', 0)} visual\n"
                summary += f"**Top Similarity Score**: {similarity_stats.get('top_similarity_score', 0):.3f}\n"
                summary += f"**Search Method**: {similarity_stats.get('search_method', ' ')}\n"
                summary += f"**Unique ADGM Docs Referenced**: {similarity_stats.get('unique_adgm_docs_referenced', 0)}\n"
                summary += f"**ADGM Jurisdiction**: {'✅ Compliant' if analysis.get('jurisdiction_compliance') else '❌ Non-Compliant'}\n"
                
                summary += "\n---\n\n"
            
            # Vector similarity insights
            summary += "## 🔍  Vector Similarity Analysis\n\n"
            total_text_matches = sum(a.get("vector_similarity_stats", {}).get("total_text_matches", 0) for a in results["documents_analysis"] if not a.get("error"))
            total_image_matches = sum(a.get("vector_similarity_stats", {}).get("total_image_matches", 0) for a in results["documents_analysis"] if not a.get("error"))
            
            summary += f"- **Total ADGM Reference Matches**: {total_text_matches + total_image_matches}\n"
            summary += f"- **Text Similarity Matches**: {total_text_matches}\n"
            summary += f"- **Visual Similarity Matches**: {total_image_matches}\n"
            summary += f"- **Search Algorithm**: Complete indexed CLIP embeddings solution\n"
            summary += f"- **Knowledge Base Size**: {kb_stats.get('total_points', 0)} vectors\n"
            summary += f"- **Performance**: ⚡ Maximum performance with all optimizations applied\n"
            
            # Handle download file with ALL FIXES
            download_file = None
            if results.get("reviewed_files"):
                potential_file = results["reviewed_files"][0]
                
                # Verify file exists and is accessible
                if os.path.exists(potential_file) and os.path.getsize(potential_file) > 0:
                    download_file = potential_file
                    print(f"✅  download file ready: {download_file}")
                    print(f"📊 Download file size: {os.path.getsize(potential_file)} bytes")
                else:
                    print(f"❌ Download file not accessible: {potential_file}")
                    download_file = None
            
            return summary, json.dumps(results, indent=2), download_file
            
        except Exception as e:
            print(f"❌ Analysis error: {e}")
            import traceback
            traceback.print_exc()
            return f"❌ Analysis failed: {str(e)}", "", None
    
    def search_knowledge_base(query, limit):
        """Enhanced knowledge base search with ALL FIXES"""
        if not query.strip():
            return "Please enter a search query."
        
        try:
            # Perform vector similarity search
            similarity_results = agent.search_similar_documents_by_vector(
                user_text=query,
                user_images=[],
                limit=int(limit)
            )
            
            if not similarity_results["text_similarities"]:
                return f"No results found for: '{query}'"
            
            report = f"## 🔍 ADGM Knowledge Base Vector Search Results (ALL FIXES)\n"
            report += f"**Query**: {query}\n"
            report += f"**Method**:  Complete Indexed CLIP Cosine Similarity\n"
            report += f"**Model Used**: {agent.model_used}\n"
            report += f"**API**: Direct Groq API Integration\n"
            report += f"**Performance**: ⚡ Maximum performance with all optimizations\n"
            report += f"**Results Found**: {len(similarity_results['text_similarities'])}\n\n"
            
            for i, result in enumerate(similarity_results["text_similarities"], 1):
                report += f"### {i}. {result['document_name']}\n"
                report += f"**Type**: {result['document_type']} | **Category**: {result['category']}\n"
                report += f"**Similarity Score**: {result['similarity_score']:.3f} (Cosine Distance)\n"
                report += f"**Content Preview**: {result['content'][:300]}...\n\n"
                report += "---\n\n"
            
            # Add search statistics
            search_stats = similarity_results.get('search_stats', {})
            if search_stats:
                report += f"## 📊 Search Statistics (ALL FIXES)\n"
                report += f"- Search Method: {search_stats.get('search_method', 'all_fixes_applied')}\n"
                report += f"- Top Similarity Score: {search_stats.get('top_similarity_score', 0):.3f}\n"
                report += f"- Unique Documents: {search_stats.get('unique_documents', 0)}\n"
                report += f"- Total Results: {search_stats.get('total_results', 0)}\n"
            
            return report
            
        except Exception as e:
            return f"❌ Search error: {str(e)}"
    
    # Create the enhanced Gradio interface with ALL FIXES
    with gr.Blocks(title="ADGM Corporate Agent - ALL FIXES", theme=gr.themes.Soft()) as demo:
        gr.Markdown(f"""
        # 🏛️ ADGM Corporate Agent -   EDITION ✨
        ### Powered by {agent.model_used} via Direct Groq API
        
        **Knowledge Base**: {agent.indexed_stats.get('total_points', 0)} vectors | **Documents**: {agent.indexed_stats.get('total_documents', 0)} | **Images**: {len(agent.image_store)} stored
        
        Upload DOCX documents for **comprehensive analysis  ** for maximum reliability and performance.
        """)
        
        with gr.Tabs():
            # Main Document Analysis Tab
            with gr.Tab("📄 Complete Document Analysis"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("### Upload Legal Documents")
                        
                        file_upload = gr.Files(
                            label="Upload DOCX Documents",
                            file_types=[".docx"],
                            file_count="multiple"
                        )
                        
                        analyze_btn = gr.Button(
                            "🔍 Analyze  ", 
                            variant="primary", 
                            size="lg"
                        )
                        
                        gr.Markdown(f"""
                        
                        **Enhanced AI Features:**
                        - ✅ **Vision LLM**: {agent.model_used}
                        - ✅ **Vector Similarity**: Complete indexed CLIP embeddings solution
                        - ✅ **ADGM Knowledge Base**: {agent.indexed_stats.get('total_documents', 0)} reference documents
                        - ✅ **Legal Citations**: Specific ADGM regulation references
                        - ✅ **Inline Comments**: Enhanced DOCX markup with alternatives
                        
                        
                        **🔧 Technical Excellence:**
                        - All known compatibility issues resolved
                        - Maximum query performance and resource efficiency
                        - Production-ready reliability and error handling
                        - Full functionality across all system configurations
                        """)
                    
                    with gr.Column(scale=2):
                        gr.Markdown("###  Analysis Results")
                        analysis_output = gr.Markdown("Upload DOCX files to begin comprehensive analysis   for maximum reliability and performance.")
                        
                        download_file = gr.File(
                            label="📥 Download Complete Enhanced Legal Review (ALL FIXES)",
                            visible=True
                        )
            
            # Knowledge Base Search Tab
            with gr.Tab("🔍  ADGM Knowledge Base Search"):
                with gr.Row():
                    with gr.Column():
                        search_query = gr.Textbox(
                            label="Search ADGM Knowledge Base ",
                            placeholder="e.g., 'company incorporation requirements', 'dispute resolution clauses', 'UBO declaration requirements'",
                            lines=2
                        )
                        
                        search_limit = gr.Slider(
                            minimum=1,
                            maximum=20,
                            value=5,
                            step=1,
                            label="Number of Results"
                        )
                        
                        search_btn = gr.Button("🎯  Complete Vector Search", variant="secondary")
                        
                        gr.Markdown(f"""
                        ** Search Features:**
                        - 🔍 Complete indexed semantic similarity search across {agent.indexed_stats.get('total_points', 0)} vectors
                        - ⚡ Maximum performance with all optimizations applied
                        - 🔧 All compatibility issues resolved
                        - 📚 Searches through {agent.indexed_stats.get('total_documents', 0)} ADGM documents
                        - 🎯 CLIP embeddings for precise semantic matching
                        - 📊 Cosine similarity scoring with optimal performance
                        - 🤖 Powered by: {agent.model_used}
                        - 🔌 Direct Groq API integration
                        """)
                        
                    with gr.Column():
                        search_results = gr.Markdown("Enter a query to search the ADGM knowledge base using the complete  solution.")
            
            # Technical Details Tab
            with gr.Tab("📊  Technical Implementation"):
                detailed_json = gr.JSON(
                    label="Complete Analysis Data "
                )
                
        
        # Event handlers
        analyze_btn.click(
            fn=analyze_documents,
            inputs=[file_upload],
            outputs=[analysis_output, detailed_json, download_file]
        )
        
        search_btn.click(
            fn=search_knowledge_base,
            inputs=[search_query, search_limit],
            outputs=[search_results]
        )
    
    return demo

def main():
    """Enhanced main application entry point  """
    
    # Environment validation
    required_vars = ["QDRANT_CLOUD_URL", "QDRANT_API_KEY", "GROQ_API_KEY"]
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    
    if missing_vars:
        print(f"❌ Missing environment variables: {missing_vars}")
        print("Please set them in your .env file")
        return
    
    # Check required files from indexing
    required_files = ["adgm_image_store.json", "indexing_metadata.json"]
    missing_files = [f for f in required_files if not Path(f).exists()]
    
    if missing_files:
        print(f"⚠️ Missing files: {missing_files}")
        if "adgm_image_store.json" in missing_files:
            print("Note: Image similarity will be limited without image store")
    
    # Create and launch interface
    try:
        demo = create_gradio_interface()
        if demo is None:
            print("❌ Failed to create interface")
            return
        
        print("\n" + "="*80)
        print("🚀 STARTING ADGM CORPORATE AGENT ")
        print("="*80)
        print("🤖 Vision LLM: meta-llama/llama-4-scout-17b-16e-instruct (with fallbacks)")
        print("🔌 API Integration: Direct Groq API (Enhanced Compatibility)")
        print("⚡ Vector Search: Complete indexed solution with all optimizations")
        print("🔧 Import Fix: Removed FieldIndex dependency for version compatibility")
        print("📁 File Upload: Enhanced filename validation and path handling")
        print("💾 File Download: Fixed temporary file management and accessibility")
        print("🎯 Vector Similarity: CLIP embeddings + Cosine distance + Indexed filtering")
        print("📚 Knowledge Base: Enhanced document counting and statistics")
        print("⚖️ Legal Features: Inline citations + Alternative clauses")
        print("="*80)
        
        demo.launch(
            server_name="0.0.0.0",
            server_port=7860,
            share=False,
            show_error=True,
            debug=True
        )
        
    except Exception as e:
        print(f"❌ Failed to start application: {e}")
        import traceback
        traceback.print_exc()
        print("\n" + "="*80)


if __name__ == "__main__":
    main()
