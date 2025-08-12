"""
One-time document indexing script for ADGM Corporate Agent
Uses meta-llama/llama-4-scout-17b-16e-instruct (Vision LLM)
Fixed Point IDs + Complete Image Data Store + Vector Similarity
"""

import pandas as pd
from pathlib import Path
import os
from dotenv import load_dotenv
import fitz  # PyMuPDF
import docx
from transformers import CLIPProcessor, CLIPModel
from PIL import Image
import torch
import numpy as np
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct
from langchain.text_splitter import RecursiveCharacterTextSplitter
import base64
import io
import uuid
import json
from typing import List, Dict, Optional
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from collections import Counter
from datetime import datetime
import shutil
import warnings
warnings.filterwarnings("ignore")

load_dotenv()

class OneTimeDocumentIndexer:
    """Complete ADGM Document Indexer with Image Storage and Fixed Point IDs"""
    
    def __init__(self, qdrant_url: str, qdrant_api_key: str, csv_path: str = "crawl_results.csv"):
        # Initialize Qdrant Cloud
        self.qdrant_client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key)
        self.collection_name = "adgm_reference_docs"
        self.csv_path = csv_path
        
        # Initialize CLIP model for unified embeddings
        print("🔄 Loading CLIP model for unified text/image embeddings...")
        self.clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
        self.clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
        self.clip_model.eval()
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        
        # Image data store - COMPLETE IMPLEMENTATION
        self.image_store_file = "adgm_image_store.json"
        self.image_store = self._load_or_create_image_store()
        self.image_metadata_file = "adgm_image_metadata.json"
        self.image_metadata = self._load_or_create_image_metadata()
        
        # Statistics tracking
        self.processing_stats = {
            "total_images_processed": 0,
            "total_images_stored": 0,
            "total_images_failed": 0,
            "storage_size_mb": 0,
            "image_types": {}
        }
        
        self._setup_qdrant_collection()
    
    def _generate_valid_point_id(self) -> str:
        """Generate valid Qdrant Point ID (FIXED - UUID only)"""
        return str(uuid.uuid4())
    
    def _load_or_create_image_store(self) -> Dict:
        """Load existing image store or create new one"""
        if Path(self.image_store_file).exists():
            try:
                with open(self.image_store_file, 'r') as f:
                    store = json.load(f)
                    print(f"📷 Loaded existing image store with {len(store)} images")
                    return store
            except Exception as e:
                print(f"Error loading image store: {e}")
        
        print("🆕 Creating new image store")
        return {}
    
    def _load_or_create_image_metadata(self) -> Dict:
        """Load or create image metadata tracking"""
        if Path(self.image_metadata_file).exists():
            try:
                with open(self.image_metadata_file, 'r') as f:
                    metadata = json.load(f)
                    print(f"📊 Loaded image metadata: {len(metadata)} entries")
                    return metadata
            except Exception as e:
                print(f"Error loading image metadata: {e}")
        
        return {
            "created_timestamp": datetime.now().isoformat(),
            "total_images": 0,
            "images_by_type": {},
            "images_by_category": {},
            "average_size_kb": 0,
            "last_updated": datetime.now().isoformat()
        }
    
    def _save_image_store(self):
        """Save image store to file with compression"""
        try:
            with open(self.image_store_file, 'w') as f:
                json.dump(self.image_store, f, indent=2)
            
            # Calculate storage statistics
            total_size = sum(len(img_data.get("base64_data", "")) for img_data in self.image_store.values())
            self.processing_stats["storage_size_mb"] = total_size / (1024 * 1024) * 0.75  # Approximate decoded size
            
            print(f"💾 Saved {len(self.image_store)} images to {self.image_store_file}")
            print(f"📦 Storage size: {self.processing_stats['storage_size_mb']:.2f} MB")
        except Exception as e:
            print(f"Error saving image store: {e}")
    
    def _save_image_metadata(self):
        """Save image metadata"""
        try:
            self.image_metadata.update({
                "total_images": len(self.image_store),
                "last_updated": datetime.now().isoformat(),
                "processing_stats": self.processing_stats
            })
            
            with open(self.image_metadata_file, 'w') as f:
                json.dump(self.image_metadata, f, indent=2, default=str)
            
            print(f"📈 Saved image metadata with {self.image_metadata['total_images']} total images")
        except Exception as e:
            print(f"Error saving image metadata: {e}")
    
    def _setup_qdrant_collection(self):
        """Setup Qdrant collection - recreate if exists (FIXED Point IDs)"""
        try:
            collections = self.qdrant_client.get_collections()
            existing_names = [col.name for col in collections.collections]
            
            if self.collection_name in existing_names:
                print(f"🗑️ Deleting existing collection: {self.collection_name}")
                self.qdrant_client.delete_collection(self.collection_name)
            
            print(f"🔧 Creating fresh collection: {self.collection_name}")
            self.qdrant_client.create_collection(
                collection_name=self.collection_name,
                vectors_config=VectorParams(size=512, distance=Distance.COSINE)
            )
            print(f"✅ Collection created successfully with COSINE distance")
            
        except Exception as e:
            print(f"❌ Error setting up collection: {e}")
            raise
    
    def embed_text(self, text: str) -> np.ndarray:
        """Generate CLIP text embedding for vector similarity"""
        inputs = self.clip_processor(
            text=text, return_tensors="pt", padding=True, 
            truncation=True, max_length=77
        )
        with torch.no_grad():
            features = self.clip_model.get_text_features(**inputs)
            features = features / features.norm(dim=-1, keepdim=True)
            return features.squeeze().numpy()
    
    def embed_image(self, image: Image.Image) -> np.ndarray:
        """Generate CLIP image embedding for vector similarity"""
        inputs = self.clip_processor(images=image, return_tensors="pt")
        with torch.no_grad():
            features = self.clip_model.get_image_features(**inputs)
            features = features / features.norm(dim=-1, keepdim=True)
            return features.squeeze().numpy()
    
    def _store_image_with_metadata(self, image: Image.Image, image_id: str, metadata: Dict) -> bool:
        """Store image with comprehensive metadata"""
        try:
            # Convert image to base64
            buffered = io.BytesIO()
            # Optimize image before storing
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Resize if too large to save storage
            max_size = (1024, 1024)
            if image.size[0] > max_size[0] or image.size[1] > max_size[1]:
                image.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            image.save(buffered, format="PNG", optimize=True)
            image_base64 = base64.b64encode(buffered.getvalue()).decode()
            
            # Calculate size
            size_kb = len(image_base64) / 1024
            
            # Store with comprehensive metadata
            self.image_store[image_id] = {
                "base64_data": image_base64,
                "image_type": metadata.get("image_type", "unknown"),
                "source_file": metadata.get("source_file", ""),
                "document_id": metadata.get("document_id", ""),
                "document_name": metadata.get("document_name", "Unknown"),
                "category": metadata.get("category", "General"),
                "document_type": metadata.get("document_type", "Unknown"),
                "created_timestamp": datetime.now().isoformat(),
                "image_dimensions": f"{image.width}x{image.height}",
                "size_kb": round(size_kb, 2),
                "compression_applied": image.size != (image.width, image.height),
                **metadata  # Include all other metadata
            }
            
            # Update statistics
            self.processing_stats["total_images_stored"] += 1
            img_type = metadata.get("image_type", "unknown")
            self.processing_stats["image_types"][img_type] = self.processing_stats["image_types"].get(img_type, 0) + 1
            
            return True
            
        except Exception as e:
            print(f"❌ Error storing image {image_id}: {e}")
            self.processing_stats["total_images_failed"] += 1
            return False
    
    def _process_pdf_with_fixed_ids(self, pdf_path: str, metadata: Dict) -> Dict:
        """Process PDF document with FIXED Point IDs and image storage"""
        stats = {"text_chunks": 0, "images": 0, "errors": []}
        
        try:
            doc = fitz.open(pdf_path)
            document_id = str(uuid.uuid4())
            points_to_insert = []
            
            for page_num, page in enumerate(doc):
                # Process text with FIXED Point IDs
                text = page.get_text()
                if text.strip():
                    chunks = self.text_splitter.split_text(text)
                    for chunk_idx, chunk in enumerate(chunks):
                        embedding = self.embed_text(chunk)
                        point = PointStruct(
                            id=self._generate_valid_point_id(),  # ✅ FIXED: Pure UUID
                            vector=embedding.tolist(),
                            payload={
                                "document_id": document_id,
                                "content": chunk,
                                "content_type": "text",
                                "page_number": page_num,
                                "chunk_index": chunk_idx,  # Moved to payload
                                "source_file": pdf_path,
                                **metadata
                            }
                        )
                        points_to_insert.append(point)
                        stats["text_chunks"] += 1
                
                # Process images with storage
                for img_idx, img in enumerate(page.get_images(full=True)):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        pil_image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                        
                        # Create image ID and metadata
                        image_id = f"{document_id}_pdf_p{page_num}_i{img_idx}"
                        img_metadata = {
                            "image_type": "pdf_embedded",
                            "source_file": pdf_path,
                            "document_id": document_id,
                            "page_number": page_num,
                            "image_index": img_idx,
                            **metadata
                        }
                        
                        # Store image with metadata
                        if self._store_image_with_metadata(pil_image, image_id, img_metadata):
                            # Create embedding for vector search
                            embedding = self.embed_image(pil_image)
                            point = PointStruct(
                                id=self._generate_valid_point_id(),  # ✅ FIXED: Pure UUID
                                vector=embedding.tolist(),
                                payload={
                                    "document_id": document_id,
                                    "content": f"[Image from page {page_num}]",
                                    "content_type": "image",
                                    "page_number": page_num,
                                    "image_index": img_idx,  # Moved to payload
                                    "image_id": image_id,  # Reference to image store
                                    "source_file": pdf_path,
                                    **metadata
                                }
                            )
                            points_to_insert.append(point)
                            stats["images"] += 1
                        
                        self.processing_stats["total_images_processed"] += 1
                        
                    except Exception as e:
                        stats["errors"].append(f"Image error in page {page_num}: {e}")
                        self.processing_stats["total_images_failed"] += 1
            
            doc.close()
            
            # Batch insert to Qdrant
            if points_to_insert:
                self.qdrant_client.upsert(
                    collection_name=self.collection_name,
                    points=points_to_insert
                )
            
        except Exception as e:
            stats["errors"].append(f"PDF processing error: {e}")
        
        return stats
    
    def _process_docx_with_fixed_ids_and_storage(self, docx_path: str, metadata: Dict) -> Dict:
        """Process DOCX with FIXED Point IDs and complete image storage"""
        stats = {"text_chunks": 0, "images": 0, "errors": []}
        
        try:
            doc = docx.Document(docx_path)
            document_id = str(uuid.uuid4())
            points_to_insert = []
            
            # Process text content with FIXED Point IDs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            document_text = '\n'.join(full_text)
            
            if document_text.strip():
                chunks = self.text_splitter.split_text(document_text)
                for chunk_idx, chunk in enumerate(chunks):
                    embedding = self.embed_text(chunk)
                    point = PointStruct(
                        id=self._generate_valid_point_id(),  # ✅ FIXED: Pure UUID
                        vector=embedding.tolist(),
                        payload={
                            "document_id": document_id,
                            "content": chunk,
                            "content_type": "text",
                            "chunk_index": chunk_idx,  # Moved to payload
                            "source_file": docx_path,
                            **metadata
                        }
                    )
                    points_to_insert.append(point)
                    stats["text_chunks"] += 1
            
            # Process tables as images with complete storage
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    if table_data and any(any(cell for cell in row) for row in table_data):
                        # Create table image
                        table_image = self._create_table_image_enhanced(table_data)
                        
                        if table_image:
                            # Create image ID and comprehensive metadata
                            image_id = f"{document_id}_table_{table_idx}"
                            table_metadata = {
                                "image_type": "table",
                                "source_file": docx_path,
                                "document_id": document_id,
                                "table_index": table_idx,
                                "table_dimensions": f"{len(table_data)}x{len(table_data[0]) if table_data else 0}",
                                "table_text_preview": str(table_data)[:200] + "..." if len(str(table_data)) > 200 else str(table_data),
                                "table_has_header": len(table_data) > 1,
                                **metadata
                            }
                            
                            # Store image with metadata
                            if self._store_image_with_metadata(table_image, image_id, table_metadata):
                                # Create embedding for vector search
                                embedding = self.embed_image(table_image)
                                point = PointStruct(
                                    id=self._generate_valid_point_id(),  # ✅ FIXED: Pure UUID
                                    vector=embedding.tolist(),
                                    payload={
                                        "document_id": document_id,
                                        "content": f"[Table from {metadata.get('document_name', 'document')}]",
                                        "content_type": "image",
                                        "image_type": "table",
                                        "table_index": table_idx,  # Moved to payload
                                        "image_id": image_id,  # Reference to image store
                                        "source_file": docx_path,
                                        **metadata
                                    }
                                )
                                points_to_insert.append(point)
                                stats["images"] += 1
                            
                            self.processing_stats["total_images_processed"] += 1
                            
                except Exception as e:
                    stats["errors"].append(f"Table processing error: {e}")
                    self.processing_stats["total_images_failed"] += 1
            
            # Process embedded images with complete storage
            try:
                for rel_idx, rel in enumerate(doc.part.rels.values()):
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            image = Image.open(io.BytesIO(image_data))
                            if image.mode != 'RGB':
                                image = image.convert('RGB')
                            
                            # Create image ID and metadata
                            image_id = f"{document_id}_embedded_{rel_idx}"
                            embedded_metadata = {
                                "image_type": "embedded",
                                "source_file": docx_path,
                                "document_id": document_id,
                                "image_index": rel_idx,
                                "original_format": image.format or "Unknown",
                                **metadata
                            }
                            
                            # Store image with metadata
                            if self._store_image_with_metadata(image, image_id, embedded_metadata):
                                # Create embedding for vector search
                                embedding = self.embed_image(image)
                                point = PointStruct(
                                    id=self._generate_valid_point_id(),  # ✅ FIXED: Pure UUID
                                    vector=embedding.tolist(),
                                    payload={
                                        "document_id": document_id,
                                        "content": f"[Embedded Image from {metadata.get('document_name', 'document')}]",
                                        "content_type": "image",
                                        "image_type": "embedded",
                                        "image_index": rel_idx,  # Moved to payload
                                        "image_id": image_id,  # Reference to image store
                                        "source_file": docx_path,
                                        **metadata
                                    }
                                )
                                points_to_insert.append(point)
                                stats["images"] += 1
                            
                            self.processing_stats["total_images_processed"] += 1
                            
                        except Exception as e:
                            stats["errors"].append(f"Embedded image error: {e}")
                            self.processing_stats["total_images_failed"] += 1
            except Exception as e:
                stats["errors"].append(f"Error accessing embedded images: {e}")
            
            # Insert to Qdrant
            if points_to_insert:
                self.qdrant_client.upsert(
                    collection_name=self.collection_name,
                    points=points_to_insert
                )
            
            # Save image store periodically
            if self.processing_stats["total_images_processed"] % 50 == 0:
                self._save_image_store()
                self._save_image_metadata()
            
        except Exception as e:
            stats["errors"].append(f"DOCX processing error: {e}")
        
        return stats
    
    def _create_table_image_enhanced(self, table_data: List[List[str]]) -> Optional[Image.Image]:
        """Create enhanced table image for Vision LLM"""
        try:
            if not table_data or not table_data[0]:
                return None
            
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            
            # Dynamic sizing based on content
            fig_width = max(cols * 2.5, 8)
            fig_height = max(rows * 0.8, 4)
            
            fig, ax = plt.subplots(figsize=(fig_width, fig_height))
            ax.set_xlim(0, cols)
            ax.set_ylim(0, rows)
            ax.axis('off')
            
            # Enhanced table rendering
            for i, row in enumerate(table_data):
                for j in range(cols):
                    cell_text = row[j] if j < len(row) else ""
                    
                    # Enhanced cell styling
                    is_header = i == 0
                    rect = patches.Rectangle(
                        (j, rows - i - 1), 1, 1,
                        linewidth=2 if is_header else 1, 
                        edgecolor='black', 
                        facecolor='lightblue' if is_header else 'white',
                        alpha=0.8 if is_header else 1.0
                    )
                    ax.add_patch(rect)
                    
                    # Enhanced text formatting
                    cell_text = str(cell_text)[:120]  # Increased limit
                    font_size = 11 if is_header else 9
                    font_weight = 'bold' if is_header else 'normal'
                    
                    ax.text(j + 0.5, rows - i - 0.5, cell_text,
                           ha='center', va='center', 
                           fontsize=font_size, 
                           weight=font_weight,
                           wrap=True)
            
            # High-quality image output
            fig.canvas.draw()
            buf = io.BytesIO()
            plt.savefig(buf, format='png', bbox_inches='tight', dpi=200, 
                       facecolor='white', edgecolor='none', 
                       pad_inches=0.1)
            buf.seek(0)
            image = Image.open(buf)
            plt.close(fig)
            
            return image
            
        except Exception as e:
            print(f"Error creating enhanced table image: {e}")
            return None
    
    def index_all_documents_from_csv(self):
        """Main indexing function with FIXED Point IDs and complete image storage"""
        try:
            if not Path(self.csv_path).exists():
                print(f"❌ CSV file not found: {self.csv_path}")
                return
            
            df = pd.read_csv(self.csv_path)
            print(f"📊 Loaded CSV with {len(df)} total entries")
            
            # Filter for successful downloads
            document_files = df[
                (df['Status'] == '✅ Success') & 
                (df['SavedLocalPath'].notna()) &
                (df['SavedLocalPath'].str.contains(r'\.(pdf|docx?|doc)$', case=False, na=False, regex=True))
            ].copy()
            
            print(f"🔍 Found {len(document_files)} successful document downloads")
            
            # Statistics tracking
            total_stats = {
                "documents_processed": 0,
                "total_text_chunks": 0,
                "total_images": 0,
                "pdf_count": 0,
                "docx_count": 0,
                "doc_count": 0,
                "errors": [],
                "missing_files": [],
                "point_id_format": "uuid_only_fixed"
            }
            
            # Process each document
            for idx, row in document_files.iterrows():
                file_path = str(row['SavedLocalPath']).replace('&amp;', '&')
                
                if not Path(file_path).exists():
                    total_stats["missing_files"].append(file_path)
                    continue
                
                file_ext = Path(file_path).suffix.lower()
                metadata = {
                    "document_name": Path(file_path).stem.replace('_', ' '),
                    "document_type": row.get('DocumentType', 'Unknown'),
                    "category": row.get('CategoryFolder', 'General'),
                    "jurisdiction": "ADGM",
                    "source_url": row.get('PageURL', ''),
                    "file_url": row.get('FileURL', ''),
                    "csv_index": idx,
                    "file_extension": file_ext,
                    "is_template": "template" in file_path.lower() or "form" in file_path.lower()
                }
                
                try:
                    print(f"🔄 Processing ({total_stats['documents_processed']+1}/{len(document_files)}): {metadata['document_name']}")
                    
                    if file_ext == '.pdf':
                        stats = self._process_pdf_with_fixed_ids(file_path, metadata)
                        total_stats["pdf_count"] += 1
                    elif file_ext in ['.docx', '.doc']:
                        stats = self._process_docx_with_fixed_ids_and_storage(file_path, metadata)
                        if file_ext == '.docx':
                            total_stats["docx_count"] += 1
                        else:
                            total_stats["doc_count"] += 1
                    
                    # Aggregate stats
                    total_stats["documents_processed"] += 1
                    total_stats["total_text_chunks"] += stats.get("text_chunks", 0)
                    total_stats["total_images"] += stats.get("images", 0)
                    total_stats["errors"].extend(stats.get("errors", []))
                    
                    # Progress updates
                    if total_stats["documents_processed"] % 10 == 0:
                        print(f"📈 Progress: {total_stats['documents_processed']}/{len(document_files)} documents processed")
                        print(f"📷 Images stored: {len(self.image_store)}")
                
                except Exception as e:
                    error_msg = f"Error processing {file_path}: {str(e)}"
                    total_stats["errors"].append(error_msg)
                    print(f"❌ {error_msg}")
            
            # Final saves
            self._save_image_store()
            self._save_image_metadata()
            
            # Final report
            self._generate_comprehensive_report(total_stats, len(document_files))
            
            # Save indexing metadata
            self._save_indexing_metadata(total_stats)
            
        except Exception as e:
            print(f"❌ Fatal error during indexing: {e}")
    
    def _generate_comprehensive_report(self, stats: Dict, total_docs: int):
        """Generate comprehensive indexing report"""
        print("\n" + "="*70)
        print("🎉 COMPLETE ADGM INDEXING FINISHED!")
        print("="*70)
        print(f"📊 Total Documents in CSV: {total_docs}")
        print(f"✅ Successfully Processed: {stats['documents_processed']}")
        print(f"❌ Missing Files: {len(stats['missing_files'])}")
        print(f"⚠️  Processing Errors: {len(stats['errors'])}")
        print(f"\n📄 File Types Processed:")
        print(f"   - PDF files: {stats['pdf_count']}")
        print(f"   - DOCX files: {stats['docx_count']}")
        print(f"   - DOC files: {stats['doc_count']}")
        print(f"\n📚 Content Indexed:")
        print(f"   - Text chunks: {stats['total_text_chunks']}")
        print(f"   - Images (including tables): {stats['total_images']}")
        print(f"\n🖼️ Image Storage Statistics:")
        print(f"   - Total images processed: {self.processing_stats['total_images_processed']}")
        print(f"   - Images successfully stored: {self.processing_stats['total_images_stored']}")
        print(f"   - Images failed: {self.processing_stats['total_images_failed']}")
        print(f"   - Storage size: {self.processing_stats['storage_size_mb']:.2f} MB")
        print(f"   - Image types: {dict(list(self.processing_stats['image_types'].items())[:5])}")
        print(f"\n🔍 Qdrant Configuration:")
        print(f"   - Collection: {self.collection_name}")
        print(f"   - Distance metric: COSINE (for vector similarity)")
        print(f"   - Point ID format: {stats.get('point_id_format', 'uuid_only_fixed')} ✅")
        print(f"   - Vector dimensions: 512 (CLIP embeddings)")
        print(f"\n🤖 Ready for Vision LLM:")
        print(f"   - Model: meta-llama/llama-4-scout-17b-16e-instruct")
        print(f"   - Multimodal: Text + Image embeddings")
        print(f"   - Vector similarity: CLIP + Cosine distance")
        print("="*70)
        
        if stats['errors']:
            print(f"\n⚠️  Sample Errors (showing 3 of {len(stats['errors'])}):")
            for error in stats['errors'][:3]:
                print(f"   - {error}")
    
    def _save_indexing_metadata(self, stats: Dict):
        """Save comprehensive indexing metadata"""
        metadata = {
            "indexing_timestamp": datetime.now().isoformat(),
            "collection_name": self.collection_name,
            "csv_file": self.csv_path,
            "vision_llm_model": "meta-llama/llama-4-scout-17b-16e-instruct",
            "approach": "fixed_point_ids_with_complete_image_storage",
            "point_id_format": "uuid_only_fixed",
            "vector_similarity": "clip_cosine_distance",
            "statistics": stats,
            "image_storage_stats": self.processing_stats,
            "images_stored": len(self.image_store),
            "image_store_file": self.image_store_file,
            "image_metadata_file": self.image_metadata_file
        }
        
        with open("indexing_metadata.json", "w") as f:
            json.dump(metadata, f, indent=2, default=str)
        
        print(f"💾 Complete indexing metadata saved to: indexing_metadata.json")
    
    def create_backup(self):
        """Create backup of all data files"""
        backup_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_files = [
            (self.image_store_file, f"backup_{backup_timestamp}_{self.image_store_file}"),
            (self.image_metadata_file, f"backup_{backup_timestamp}_{self.image_metadata_file}"),
            ("indexing_metadata.json", f"backup_{backup_timestamp}_indexing_metadata.json")
        ]
        
        for source, backup in backup_files:
            if Path(source).exists():
                shutil.copy2(source, backup)
                print(f"📁 Backup created: {backup}")
    
    def validate_image_store(self):
        """Validate image store integrity"""
        print("🔍 Validating image store integrity...")
        
        valid_count = 0
        invalid_count = 0
        total_size = 0
        
        for image_id, image_data in self.image_store.items():
            try:
                base64_data = image_data.get("base64_data", "")
                if base64_data:
                    decoded = base64.b64decode(base64_data)
                    img = Image.open(io.BytesIO(decoded))
                    valid_count += 1
                    total_size += len(base64_data)
                else:
                    invalid_count += 1
            except Exception as e:
                invalid_count += 1
                print(f"❌ Invalid image {image_id}: {e}")
        
        print(f"✅ Validation complete:")
        print(f"   - Valid images: {valid_count}")
        print(f"   - Invalid images: {invalid_count}")
        print(f"   - Total storage: {total_size / (1024*1024) * 0.75:.2f} MB")
        
        return invalid_count == 0

def main():
    """Main execution with comprehensive validation"""
    print("🏛️ ADGM Document Indexer - Complete Version")
    print("Features: Fixed Point IDs + Complete Image Storage + Vector Similarity")
    print("Model: meta-llama/llama-4-scout-17b-16e-instruct")
    print("="*70)
    
    # Load configuration
    QDRANT_URL = os.getenv("QDRANT_CLOUD_URL")
    QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
    CSV_PATH = "crawl_results.csv"
    
    # Validate configuration
    if not QDRANT_URL or not QDRANT_API_KEY:
        print("❌ Missing Qdrant Cloud credentials!")
        print("Please set QDRANT_CLOUD_URL and QDRANT_API_KEY in your .env file")
        return
    
    if not Path(CSV_PATH).exists():
        print(f"❌ CSV file not found: {CSV_PATH}")
        return
    
    # Test Qdrant connection
    try:
        test_client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY)
        collections = test_client.get_collections()
        print("✅ Qdrant connection verified")
    except Exception as e:
        print(f"❌ Qdrant connection failed: {e}")
        return
    
    # Analyze CSV
    try:
        df = pd.read_csv(CSV_PATH)
        document_count = len(df[
            (df['Status'] == '✅ Success') & 
            (df['SavedLocalPath'].notna()) &
            (df['SavedLocalPath'].str.contains(r'\.(pdf|docx?|doc)$', case=False, na=False, regex=True))
        ])
        
        print(f"📋 Found {document_count} documents to index from {len(df)} CSV entries")
        
        # Document type breakdown
        doc_types = df[df['Status'] == '✅ Success']['DocumentType'].value_counts()
        print(f"\n📊 Document Types Preview:")
        for doc_type, count in doc_types.head(5).items():
            print(f"   - {doc_type}: {count}")
        
    except Exception as e:
        print(f"❌ Error analyzing CSV: {e}")
        return
    
    # Confirm before proceeding
    confirm = input(f"\n🤔 Proceed with complete indexing of {document_count} documents?\n   This will create: Fixed Point IDs + Complete Image Storage + Vector Similarity\n   Continue? (y/N): ")
    
    if confirm.lower() != 'y':
        print("❌ Indexing cancelled by user")
        return
    
    # Initialize and run indexer
    try:
        indexer = OneTimeDocumentIndexer(QDRANT_URL, QDRANT_API_KEY, CSV_PATH)
        
        # Create backup of existing data
        indexer.create_backup()
        
        # Run complete indexing
        indexer.index_all_documents_from_csv()
        
        # Validate results
        indexer.validate_image_store()
        
        print("\n✅ COMPLETE SETUP FINISHED!")
        print("🚀 Next step: Run 'python main_app.py' to start the ADGM Corporate Agent")
        print("🎯 Features ready: Vector Similarity + Vision LLM + Image Storage")
        
    except Exception as e:
        print(f"❌ Fatal error during indexing: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
